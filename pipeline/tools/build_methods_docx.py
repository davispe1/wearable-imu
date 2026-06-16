"""
build_methods_docx.py — generate docs/Gait_Kinematics_Methods.docx

A refined, reader-friendly version of docs/methods.md as a Word document:
  * plain-language "In plain terms" intro for every pipeline stage,
  * key equations rendered as clean images (matplotlib mathtext),
  * real-data figures generated from the P01 outputs,
  * a dedicated Visualization chapter (what we see, how it is made, why each
    biomechanical parameter matters),
  * reuses the existing outputs/*.png (dashboard, window, placement, stickfigure).

Run:  python tools/build_methods_docx.py
"""
from __future__ import annotations
import json, os
import numpy as np

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Circle

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "outputs")
ASSETS = os.path.join(ROOT, "docs", "assets")
DOCX = os.path.join(ROOT, "docs", "Gait_Kinematics_Methods.docx")
BASE = os.path.join(OUT, "P01_S01_2minWalk")
os.makedirs(ASSETS, exist_ok=True)

ACCENT = RGBColor(0x1F, 0x47, 0x7B)   # deep blue
GREY = RGBColor(0x55, 0x55, 0x55)

# --------------------------------------------------------------------------- #
# Equation rendering (matplotlib mathtext -> tight PNG)
# --------------------------------------------------------------------------- #
def render_eq(lines, name, fontsize=17):
    """Render one or more math lines (mathtext) to a tight PNG; return its path."""
    if isinstance(lines, str):
        lines = [lines]
    n = len(lines)
    fig = plt.figure(figsize=(7.0, 0.55 * n + 0.1))
    for i, ln in enumerate(lines):
        fig.text(0.01, 1.0 - (i + 0.7) / n, f"${ln}$", fontsize=fontsize, va="center")
    path = os.path.join(ASSETS, f"eq_{name}.png")
    fig.savefig(path, dpi=200, bbox_inches="tight", pad_inches=0.12,
                facecolor="white")
    plt.close(fig)
    return path


# --------------------------------------------------------------------------- #
# Data figures
# --------------------------------------------------------------------------- #
def load_ts():
    return np.genfromtxt(BASE + "_timeseries.csv", delimiter=",", names=True)


def fig_pipeline():
    """Stage flow diagram with the raw-data contract wall."""
    stages = [
        ("Raw dataset", ".BIN inertial files + .c3d markers (read-only)", "#dfe7f0"),
        ("1  BIN decode", "page-aware Physilog parser -> acc / gyro / mag counts", "#dfe7f0"),
        ("2  Alignment", "locate trial in 69-min file via timestamp + clock skew", "#dfe7f0"),
        ("3  Extract + segment", "SI units, delimit walking bout, inter-sensor refine", "#dfe7f0"),
        ("RAW-DATA CONTRACT WALL", "adapters/geneva.py -> IMUTrial.imu  (IMU only past here)", "#f6d8a8"),
        ("4  Mag calibration", "hard/soft-iron ellipsoid + frame align", "#d8eedd"),
        ("5  Orientation fusion", "Madgwick 6-DOF (primary) & 9-DOF, per sensor", "#d8eedd"),
        ("6  Joint angles", "yaw-immune sagittal flexion + complementary filter", "#d8eedd"),
        ("7  Gait events", "foot-IMU strikes, cadence, stride stats", "#d8eedd"),
        ("Outputs", "timeseries.csv / summary.json + visualization", "#dfe7f0"),
        ("8  Validation", "marker reference, RMSE, heading arbiter (markers ONLY here)", "#f2dada"),
    ]
    fig, ax = plt.subplots(figsize=(8.2, 10.2))
    ax.set_xlim(0, 10); ax.set_ylim(0, len(stages))
    ax.axis("off")
    y = len(stages)
    for i, (title, sub, col) in enumerate(stages):
        y -= 1
        box = FancyBboxPatch((0.4, y + 0.12), 9.2, 0.76,
                             boxstyle="round,pad=0.02,rounding_size=0.08",
                             linewidth=1.2, edgecolor="#5a5a5a", facecolor=col)
        ax.add_patch(box)
        weight = "bold" if title[0].isdigit() or "WALL" in title else "bold"
        ax.text(0.65, y + 0.62, title, fontsize=11.5, fontweight=weight,
                color="#1f2d3d", va="center")
        ax.text(0.65, y + 0.30, sub, fontsize=8.8, color="#3a4654", va="center")
        if i < len(stages) - 1:
            ax.add_patch(FancyArrowPatch((5.0, y + 0.10), (5.0, y - 0.02),
                                         arrowstyle="-|>", mutation_scale=14,
                                         color="#5a5a5a", linewidth=1.4))
    ax.set_title("Gait Kinematics — pipeline data flow", fontsize=13, fontweight="bold")
    p = os.path.join(ASSETS, "fig_pipeline.png")
    fig.tight_layout(); fig.savefig(p, dpi=130, facecolor="white"); plt.close(fig)
    return p


def fig_flexion_vs_ref(d):
    """3-panel computed (line) vs optical reference (dots) over the first mocap window."""
    t = d["t_opt_s"]
    t0, t1 = 7.5, 10.6
    sel = (t >= t0) & (t <= t1)
    joints = [("ankle", "#4c72b0"), ("knee", "#55a868"), ("hip", "#8172b3")]
    fig, axes = plt.subplots(3, 1, figsize=(8.4, 7.2), sharex=True)
    for ax, (j, c) in zip(axes, joints):
        ax.plot(t[sel], d[f"{j}_deg"][sel], color=c, lw=1.8, label="computed (6-DOF IMU)")
        ref = d[f"{j}_ref_deg"]
        m = sel & ~np.isnan(ref)
        ax.plot(t[m], ref[m], ".", ms=4, color="#c44e52", label="optical reference")
        ax.set_ylabel(f"{j}\nflexion (deg)", fontsize=10)
        ax.grid(alpha=0.25)
        ax.legend(loc="upper right", fontsize=8, ncol=2)
    axes[-1].set_xlabel("time (s)", fontsize=10)
    fig.suptitle("Computed IMU flexion vs optical reference — one mocap window (P01, right leg)",
                 fontsize=12, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.97])
    p = os.path.join(ASSETS, "fig_flexion_vs_ref.png")
    fig.savefig(p, dpi=130, facecolor="white"); plt.close(fig)
    return p


def fig_foot_events(d):
    """Foot accel magnitude + detected strikes over a steady window."""
    t = d["t_opt_s"]
    t0, t1 = 40.0, 52.0
    sel = (t >= t0) & (t <= t1)
    fig, ax = plt.subplots(figsize=(8.4, 3.0))
    ax.plot(t[sel], d["foot_acc_mag"][sel], color="#444", lw=0.8)
    fs = sel & (d["foot_strike"] > 0)
    ax.plot(t[fs], d["foot_acc_mag"][fs], "v", ms=8, color="#c44e52", label="foot strike")
    ax.set_xlabel("time (s)"); ax.set_ylabel("foot |a| (m/s$^2$)")
    ax.set_title("Foot accelerometer magnitude with detected foot strikes (steady walking)",
                 fontsize=11, fontweight="bold")
    ax.grid(alpha=0.25); ax.legend(loc="upper right", fontsize=9)
    fig.tight_layout()
    p = os.path.join(ASSETS, "fig_foot_events.png")
    fig.savefig(p, dpi=130, facecolor="white"); plt.close(fig)
    return p


def fig_velocity(d):
    """Joint angular velocity over a steady window."""
    t = d["t_opt_s"]
    t0, t1 = 40.0, 48.0
    sel = (t >= t0) & (t <= t1)
    fig, ax = plt.subplots(figsize=(8.4, 3.0))
    for j, c in (("ankle", "#4c72b0"), ("knee", "#55a868"), ("hip", "#8172b3")):
        ax.plot(t[sel], d[f"{j}_vel_dps"][sel], color=c, lw=1.2, label=j)
    ax.set_xlabel("time (s)"); ax.set_ylabel("angular velocity (deg/s)")
    ax.set_title("Joint angular velocity (derivative of flexion) — steady walking",
                 fontsize=11, fontweight="bold")
    ax.grid(alpha=0.25); ax.legend(loc="upper right", fontsize=9, ncol=3)
    fig.tight_layout()
    p = os.path.join(ASSETS, "fig_velocity.png")
    fig.savefig(p, dpi=130, facecolor="white"); plt.close(fig)
    return p


def fig_angle_def():
    """Conceptual diagram: sagittal flexion as the gravity angle about the joint axis."""
    fig, ax = plt.subplots(figsize=(6.4, 4.0))
    ax.set_xlim(-1.2, 2.4); ax.set_ylim(-1.6, 1.4); ax.axis("off")
    # gravity reference (down)
    ax.add_patch(FancyArrowPatch((0, 0), (0, -1.2), arrowstyle="-|>",
                                 mutation_scale=16, color="#888", lw=2))
    ax.text(0.05, -1.25, r"$\hat{g}$  (gravity, world down)", color="#666", fontsize=10)
    # neutral segment direction g0
    ax.add_patch(FancyArrowPatch((0, 0), (0.0, -1.1), arrowstyle="-|>",
                                 mutation_scale=14, color="#4c72b0", lw=2.5))
    # rotated gravity in segment
    ang = np.radians(40)
    gx, gy = 1.1 * np.sin(ang), -1.1 * np.cos(ang)
    ax.add_patch(FancyArrowPatch((0, 0), (gx, gy), arrowstyle="-|>",
                                 mutation_scale=14, color="#c44e52", lw=2.5))
    ax.text(gx + 0.05, gy, r"$g$  (gravity in sensor now)", color="#c44e52", fontsize=10)
    ax.text(-0.95, -0.55, r"$g_0$  (neutral)", color="#4c72b0", fontsize=10)
    # arc for theta
    th = np.linspace(-np.pi / 2, -np.pi / 2 + ang, 40)
    ax.plot(0.55 * np.cos(th), 0.55 * np.sin(th), color="#333", lw=1.5)
    ax.text(0.42, -0.35, r"$\theta_s$", fontsize=13)
    # joint axis (out of plane)
    ax.add_patch(Circle((0, 0), 0.05, color="k"))
    ax.add_patch(Circle((0, 0), 0.12, fill=False, color="k", lw=1.2))
    ax.text(0.1, 0.12, r"$j$  (joint flexion axis, out of page)", fontsize=10)
    ax.set_title("Yaw-immune sagittal angle: signed rotation of gravity about the joint axis",
                 fontsize=11, fontweight="bold")
    fig.tight_layout()
    p = os.path.join(ASSETS, "fig_angle_def.png")
    fig.savefig(p, dpi=130, facecolor="white"); plt.close(fig)
    return p


# --------------------------------------------------------------------------- #
# Document helpers
# --------------------------------------------------------------------------- #
def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def para(doc, text, italic=False, size=None, color=None, bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p


def plain_terms(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("In plain terms.  ")
    r.bold = True; r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    return p


def eq_img(doc, path, width=5.2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))


def figure(doc, path, caption, width=6.2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for pp in cells[i].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(9)
    doc.add_paragraph()
    return t


# --------------------------------------------------------------------------- #
def build():
    d = load_ts()
    summary = json.load(open(BASE + "_summary.json"))
    g = summary["gait"]
    rmse = summary["validation_rmse_deg"]["6dof"]
    rw = summary["joint_rom_in_window_deg"]

    # ---- generate figures / equations
    p_pipeline = fig_pipeline()
    p_flexref = fig_flexion_vs_ref(d)
    p_foot = fig_foot_events(d)
    p_vel = fig_velocity(d)
    p_angle = fig_angle_def()

    eq_scale = render_eq([
        r"a\,[\mathrm{m/s^2}] = \mathrm{counts}\cdot 9.80665/2048",
        r"\omega\,[\mathrm{rad/s}] = \mathrm{counts}\cdot (1/16.384)\cdot \pi/180",
    ], "scale")
    eq_ncc = render_eq(r"r(k)=\frac{(s\star t_{demeaned})[k]}{\sqrt{\mathrm{var}_{win}(k)\cdot \sum (t-\bar t)^2}}", "ncc")
    eq_skew = render_eq(r"\mathrm{skew}=\mathrm{median}(\,\mathrm{loc}-\mathrm{pred}\,),\quad \mathrm{pred}= t_{capture}-t_{rtc}", "skew")
    eq_rms = render_eq([
        r"\mathrm{rms}(t)=\sqrt{\mathrm{movavg}(\|\omega\|^2,\ 0.5s)}",
        r"\mathrm{walking\ if}\quad \deg(\mathrm{rms}) > 60^\circ/s",
    ], "rms")
    eq_ellip = render_eq([
        r"b = -Q^{-1}n,\qquad k = 1 + n^{T}Q^{-1}n",
        r"A = \mathrm{sqrtm}(Q/k),\qquad m_{cal}=P\,A\,(m-b)",
    ], "ellip")
    eq_madg = render_eq([
        r"\dot q_\omega = \frac{1}{2}\, q\otimes[0,\omega]",
        r"\dot q = \dot q_\omega - \beta\,\nabla f/\|\nabla f\|,\qquad q\leftarrow \mathrm{normalize}(q+\dot q\,dt)",
    ], "madg")
    eq_theta = render_eq([
        r"g_{0\perp}=\mathrm{norm}(g_0-(g_0\!\cdot j)j),\quad g_\perp=\mathrm{norm}(g-(g\!\cdot j)j)",
        r"\theta_s=\mathrm{atan2}((g_{0\perp}\times g_\perp)\!\cdot j,\ \ g_\perp\!\cdot g_{0\perp})",
    ], "theta")
    eq_comp = render_eq([
        r"\alpha=\tau/(\tau+dt),\qquad \tau=0.3\,s",
        r"\mathrm{flex}[i]=\alpha\,(\mathrm{flex}[i\!-\!1]+\omega_{joint}[i]\,dt)+(1-\alpha)\,\mathrm{flex}_{grav}[i]",
    ], "comp")
    eq_cad = render_eq(r"\mathrm{cadence}=(1/\mathrm{mean}(\Delta t))\cdot 60\cdot 2\ \ \mathrm{steps/min}", "cad")
    eq_heading = render_eq(r"\mathrm{heading}(t)=\mathrm{unwrap}(\mathrm{atan2}(R(q)[1,k],\,R(q)[0,k]))", "heading")
    eq_rmse = render_eq(r"\mathrm{RMSE}=\min_{|lag|\leq 0.3s}\ \sqrt{\mathrm{mean}((a_{demean}-b_{demean,lag})^2)}", "rmse")

    # =====================================================================
    doc = Document()
    # base style
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    # ---- title page
    para(doc, "Gait Kinematics", size=30, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Wearable IMU Gait Kinematics — Methods & Results",
         size=16, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    para(doc, "From raw inertial bytes to validated joint angles, gait metrics, "
              "and an animated lateral mannequin.", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    doc.add_paragraph()
    para(doc, "Reference run: subject P01 · task 2minWalk · right leg · "
              "nodes RF/RS/RT/SA (foot, shank, thigh, pelvis) · "
              "joints ankle (RF–RS), knee (RS–RT), hip (RT–SA).",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, color=GREY)
    doc.add_page_break()

    # ---- 1. Overview
    add_heading(doc, "1  Overview — what the system does", 1)
    para(doc,
         "Gait Kinematics turns the raw output of body-worn inertial sensors (IMUs) into "
         "clinically meaningful walking measurements: how each leg joint flexes through "
         "time, how fast it moves, how often the foot strikes the ground, and how those "
         "results compare against a gold-standard optical motion-capture system. It runs "
         "on the Grouvel et al. (2023, Geneva) dataset — eight Physilog 6S IMUs recorded "
         "alongside optical markers and force plates — but it is built as a rehearsal for "
         "real wearable hardware: the kinematic core consumes only raw accelerometer, "
         "gyroscope, and magnetometer data, exactly as a live device would stream it.")
    para(doc,
         "Each IMU reports three things at 256 Hz: linear acceleration (which, at rest, "
         "points along gravity), angular velocity (rotation rate), and the local magnetic "
         "field. Four sensors on one leg — foot, shank, thigh, and pelvis — let us measure "
         "the three sagittal joints between them: ankle, knee, and hip.")
    figure(doc, p_pipeline, "Figure 1. The full pipeline. Everything above the orange "
           "wall is dataset-specific data ingest; everything below sees only IMU signals. "
           "Validation (red) reads optical markers, but only for scoring — never to compute "
           "the result.", width=5.6)

    add_heading(doc, "1.1  The raw-data contract", 2)
    plain_terms(doc,
        "There is a strict wall in the code. Everything that computes a joint angle or a "
        "gait event is allowed to read only the IMU signals. The optical markers are read "
        "in exactly one place — validation — and only to score accuracy. This guarantees "
        "the results would be identical on a live device with no cameras in the room.")
    para(doc,
         "The wall is enforced by adapters/geneva.py (the IMUTrial object) and proven by "
         "selftest.py, which scrambles the marker data and confirms that every computed "
         "output is bit-for-bit unchanged (Section 11).")

    # ---- 2. The data
    add_heading(doc, "2  The input data", 1)
    para(doc, "Each sensor records one continuous proprietary .BIN file for the whole "
              "~69-minute session. The raw channels decoded from it:")
    table(doc, ["Tag", "Channel", "Rate", "Units after decode", "Used for"],
          [["0x13", "Accelerometer", "256 Hz", "m/s² (÷2048 ×9.80665)", "gravity / tilt, foot impacts"],
           ["0x14", "Gyroscope", "256 Hz", "rad/s (÷16.384, °→rad)", "rotation, joint axis, events"],
           ["0x18", "Magnetometer", "256 Hz", "raw counts (calib. normalises)", "9-DOF heading (comparison)"],
           ["0x15", "Barometer", "64 Hz", "two scalars (not a vector)", "housekeeping only"]])
    para(doc, "A key finding: the magnetometer is the 256 Hz tag 0x18, sample-aligned with "
              "accel/gyro — not the 64 Hz channel the dataset documentation suggested (that "
              "is the barometer). So no up-sampling is needed; every motion sample has a "
              "co-timed magnetic reading.", italic=True)

    # ---- 3..  Stages
    add_heading(doc, "3  Stage 1 — Decoding the binary files", 1)
    plain_terms(doc,
        "The sensor files are in an undocumented format, so the first job is to reverse-"
        "engineer them into plain numbers in physical units.")
    para(doc, "The file is a sequence of 512-byte pages. Page 0 holds configuration and the "
              "real-time-clock start. Every later page is an 8-byte header (page index + "
              "cumulative sample counter) followed by 63 fixed 8-byte records, each "
              "record = tag(1) + counter(1) + three big-endian int16 values. The parser is "
              "fully vectorised and drops only ~0.002 % unknown-tag records. Counts convert "
              "to SI units as:")
    eq_img(doc, eq_scale, 5.4)
    para(doc, "Correctness is confirmed by physics (|a| ≈ 9.79 m/s² and ω ≈ 0 at rest) and "
              "by cross-correlation against the dataset's sync export. The magnetometer's "
              "absolute scale is unknown, so only its direction is used.")

    add_heading(doc, "4  Stage 2 — Finding the trial inside the session", 1)
    plain_terms(doc,
        "The walk we care about is a few minutes buried inside a 69-minute file. A short "
        "walking pattern repeats and is not unique, so we cannot find it by pattern-matching "
        "alone — we use the clocks instead.")
    para(doc, "The optical capture file stores an absolute start time; the IMU file stores "
              "its real-time-clock start. They differ by a small, near-constant session "
              "skew. We estimate the skew once from the long calibration trials by cross-"
              "correlating gyroscope magnitude, then take the median over confident anchors:")
    eq_img(doc, eq_skew, 5.0)
    para(doc, "The sliding correlation is a fast normalised (Pearson) cross-correlation:")
    eq_img(doc, eq_ncc, 4.6)
    para(doc, "For P01 the skew is −12.14 s and holds to within ±1.5 s across the whole "
              "session. Short walking templates are not used to refine timing; the "
              "timestamp + skew anchor (good to ~±0.4 s) is kept.")

    add_heading(doc, "5  Stage 3 — Extraction & segmentation", 1)
    plain_terms(doc,
        "Convert every sensor to physical units, put all eight on one shared clock, then cut "
        "out just the continuous walking bout and mark the turnarounds at the ends of the "
        "walkway.")
    para(doc, "The walking bout is delimited from the foot gyroscope: a smoothed RMS of "
              "rotation rate crossing a threshold marks walking.")
    eq_img(doc, eq_rms, 4.8)
    para(doc, "Short gaps are bridged, the run covering the optical anchors is kept and "
              "padded — for P01 a single continuous 128.7 s bout. Adjacent leg sensors are "
              "then refined to <25 ms using shared heel-strike impacts (an 8 Hz high-pass "
              "accel feature cross-correlated against the foot). The pelvis impacts are too "
              "damped to refine, so it stays on the ~±0.4 s clock — the origin of the hip's "
              "lower confidence. Turnarounds are detected on the pelvis vertical-axis "
              "rotation rate (a turn integrates to ≥120°); for P01 there are 7 turns, which "
              "are excluded from all gait statistics.")

    add_heading(doc, "6  Stage 4 — Magnetometer calibration", 1)
    plain_terms(doc,
        "A raw magnetometer is distorted by nearby metal (hard- and soft-iron effects). "
        "Calibration finds the transform that maps its readings back onto a clean sphere, "
        "and figures out how the magnetometer's axes line up with the accelerometer's.")
    para(doc, "An algebraic ellipsoid fit gives the hard-iron offset b and soft-iron matrix A:")
    eq_img(doc, eq_ellip, 5.2)
    para(doc, "The chip axes are then aligned to the accel frame by searching all 48 signed "
              "axis permutations and choosing the one that makes the magnetic dip angle most "
              "constant. For P01 the fit is poor (sphere residual 0.65–0.83) and the local "
              "inclination is ~50° versus Geneva's geomagnetic 63° — clear evidence of "
              "indoor field distortion. A separate world-frame test confirms tag 0x18 is a "
              "genuine magnetometer (its world inclination stays ~−50° across rotations up to "
              "91°), but the distorted field is exactly why the magnetometer cannot help here.")

    add_heading(doc, "7  Stage 5 — Orientation fusion (Madgwick)", 1)
    plain_terms(doc,
        "Each sensor's acceleration, rotation, and (optionally) magnetic field are fused "
        "into an orientation — a quaternion saying how the sensor is tilted in the world — "
        "at every instant. We compute a 6-DOF version (accel + gyro) as primary and a 9-DOF "
        "version (adding the magnetometer) for comparison.")
    para(doc, "The Madgwick filter blends gyro integration with a gradient-descent "
              "correction toward the direction gravity (and the magnetic field) imply:")
    eq_img(doc, eq_madg, 5.4)
    para(doc, "Gains are β = 0.033 (6-DOF) and 0.05 (9-DOF). Because the pelvis sits near the "
              "Euler yaw singularity (pitch 71–89°), heading is read not from Euler yaw but "
              "from the most-horizontal body axis:")
    eq_img(doc, eq_heading, 5.4)
    para(doc, "Fusion runs per sensor on its own clock; cross-sensor alignment happens later "
              "on the orientation outputs. 6-DOF is primary; 9-DOF is degraded by the "
              "distorted field.")

    add_heading(doc, "8  Stage 6 — Joint angles (the core result)", 1)
    plain_terms(doc,
        "A joint angle is the difference in orientation between two segments. The clever part "
        "is making it immune to heading drift: instead of using full 3D orientation (whose "
        "yaw slowly drifts), we measure how the gravity direction rotates about each "
        "segment's flexion axis. Gravity never drifts, so the angle stays clean for minutes.")
    figure(doc, p_angle, "Figure 2. Sagittal flexion is the signed angle through which "
           "gravity has rotated about the joint's flexion axis, relative to the neutral "
           "standing pose. Because it is anchored to gravity, it carries no heading drift.",
           width=4.8)
    para(doc, "The flexion axis j of each segment is its largest-variance gyro direction, "
              "estimated on steady (non-turning) samples. With gravity expressed in the "
              "sensor frame, each segment's rotation from neutral is:")
    eq_img(doc, eq_theta, 5.4)
    para(doc, "The joint flexion is the distal-minus-proximal segment rotation. Because the "
              "gravity projection lags during fast motion, a complementary filter blends it "
              "with the integrated joint rate:")
    eq_img(doc, eq_comp, 5.6)
    para(doc, "Angular velocity and acceleration are time-derivatives of flexion; an "
              "anatomical sign per joint is applied so the output matches clinical "
              "convention. The method reports sagittal flexion only.")
    figure(doc, p_flexref, "Figure 3. The computed IMU flexion (solid) tracks the optical "
           "reference (dots) closely over a motion-capture window. This is the central "
           "validation: the IMU-only result reproduces the camera-based gold standard.")

    add_heading(doc, "9  Stage 7 — Gait events", 1)
    plain_terms(doc,
        "From the foot sensor we detect each step — mid-swing, foot strike, toe-off — and "
        "from the rhythm of foot strikes we compute cadence and stride time.")
    para(doc, "Mid-swing peaks are found on the foot sagittal rotation rate; foot strike is "
              "the rate minimum just after mid-swing, toe-off the minimum just before. "
              "Cadence comes from steady-state stride intervals (one foot instrumented, so "
              "doubled for both feet):")
    eq_img(doc, eq_cad, 4.8)
    figure(doc, p_foot, "Figure 4. Foot accelerometer magnitude; each red marker is a "
           "detected foot strike. The regular spacing is the stride rhythm that yields "
           "cadence and stride-time statistics.")

    add_heading(doc, "10  Stage 8 — Validation against optical motion capture", 1)
    plain_terms(doc,
        "We score the IMU result against the camera system. The reference angles are built "
        "from the optical joint centres, zeroed at the same neutral pose, and compared after "
        "removing a constant offset and the best small time lag.")
    para(doc, "The reference uses the c3d-provided joint centres (hip/knee/ankle) and pelvis "
              "markers; each joint is the signed angle in the plane perpendicular to the "
              "pelvis mediolateral axis. The error metric is offset-removed and lag-optimised:")
    eq_img(doc, eq_rmse, 5.6)
    para(doc, "A heading arbiter compares IMU pelvis heading (6-DOF and 9-DOF) to the "
              "optical heading; foot strikes are checked against the dataset's Zeni events. "
              "Headline results (in-window, computed vs optical) appear in Section 12.")

    add_heading(doc, "11  Stage 9 — The raw-data contract proof", 1)
    plain_terms(doc,
        "A regression test that proves the core never cheats by peeking at the cameras.")
    para(doc, "selftest.py runs the core, then time-shuffles the marker frames and drops the "
              "event labels and runs it again. It asserts that 21 invariants (flexion, "
              "angular velocity, ROM for each joint × fusion mode, plus foot strikes, "
              "cadence, stride count) are bit-for-bit identical, while the validation RMSE "
              "changes and matched steps drop to zero. PASS ⇒ the core is a pure function of "
              "the IMU signals.")

    # ---- 12 Visualization (the big new chapter)
    doc.add_page_break()
    add_heading(doc, "12  Visualization — reading the results", 1)
    para(doc,
         "The visualization is where the numbers become understandable. visualize.py opens a "
         "single interactive window (or renders the same content to PNG stills with --save). "
         "It deliberately shows the same data three ways at once — as an animated body, as "
         "synchronised time-series plots, and as a metrics summary — linked by one moving "
         "time-cursor so a feature in a curve can be seen in the moving limb at the same "
         "instant.")
    figure(doc, BASE + "_window.png",
           "Figure 5. The interactive window. Left: the lateral mannequin (top) driven by "
           "the clean sagittal flexion angles, and the static sensor-placement diagram "
           "(bottom). Right: synchronised joint-flexion, angular-velocity, and foot-impact "
           "panels with a sweeping time-cursor, plus a live metrics box.")

    add_heading(doc, "12.1  The animated lateral mannequin", 2)
    para(doc,
         "A side-view (sagittal) stick body — thigh, shank, foot, and a pelvis/trunk stub — "
         "is reconstructed each frame purely from the three computed flexion angles and the "
         "subject's segment lengths (scaled from stature). Crucially it is driven by the "
         "clean yaw-immune flexion angles, not the full 3D orientation: the earlier 3D view "
         "looked wrong because its yaw drifted, whereas the gravity-anchored flexion is "
         "stable. Playback advances by real wall-clock time mapped through the 256 Hz data, "
         "so 1.0× is true real-time regardless of how fast the screen redraws. Optional "
         "forward translation walks the figure across the view at cadence × step length; the "
         "root is otherwise fixed because the pipeline computes orientation, not global "
         "position.")
    para(doc, "Why it matters: it is an immediate sanity check that the angles are "
              "anatomically plausible — the knee bends the right way through swing, the "
              "ankle dorsi/plantarflexes in phase with the foot — something a table of "
              "numbers cannot convey.")

    add_heading(doc, "12.2  The synchronised gait panels", 2)
    para(doc, "Three stacked, time-linked panels share the cursor with the mannequin:")
    para(doc, "• Joint flexion (deg) — ankle, knee, hip versus time (Figure 3 shows these "
              "against the optical reference). The shape and amplitude are the primary "
              "kinematic result.")
    para(doc, "• Joint angular velocity (deg/s) — the time-derivative of flexion. Peaks mark "
              "the fastest joint motion (e.g. knee swing); it is what reveals dynamic, not "
              "just static, range.")
    eq_img(doc, render_eq(r"\omega_{joint}(t)=\frac{d}{dt}\,\mathrm{flexion}(t)", "veldef"), 3.4)
    figure(doc, p_vel, "Figure 6. Joint angular velocity over a few strides — the knee and "
           "ankle reach the highest rates during swing.")
    para(doc, "• Foot acceleration magnitude + strikes — the impact signal with detected "
              "foot strikes (Figure 4), the basis of cadence and stride timing.")

    add_heading(doc, "12.3  Sensor placement & stick-figure stills", 2)
    figure(doc, BASE + "_placement.png",
           "Figure 7. Sensor placement: the four IMU nodes — RF (foot), RS (shank), "
           "RT (thigh), SA (pelvis/sacrum) — on a neutral-stance lateral body. This is the "
           "physical setup every measurement derives from.", width=3.6)
    figure(doc, BASE + "_stickfigure.png",
           "Figure 8. The right-leg sagittal chain (pelvis→hip→knee→ankle→toe) drawn at "
           "several frames across one steady stride, coloured dark→yellow in time. The "
           "evolving knee bend and foot clearance through swing are visible at a glance.")

    add_heading(doc, "12.4  Why these are the right biomechanical parameters", 2)
    para(doc, "Every quantity shown is a standard, clinically interpretable gait measure:")
    table(doc, ["Parameter", "What it is", "Why it matters"],
          [["Sagittal flexion / ROM", "joint angle and its range through the stride",
            "the core of gait assessment — reduced knee or ankle ROM flags pathology, "
            "recovery, or compensation"],
           ["Angular velocity", "how fast each joint moves",
            "distinguishes fast vs slow/guarded movement; peak swing velocity is a "
            "sensitive marker even when ROM looks normal"],
           ["Cadence", "steps per minute",
            "a primary descriptor of walking speed and rhythm, and an outcome measure in "
            "rehabilitation"],
           ["Stride time (mean ± SD)", "duration and variability of the gait cycle",
            "variability is a recognised marker of stability and fall risk"],
           ["Foot strikes / steady strides", "step timing events",
            "segment the gait cycle and gate which strides are 'clean' straight walking"],
           ["Turnarounds", "walkway end-turns",
            "excluded so metrics reflect steady gait, not turning"]])
    para(doc, "Because each panel is time-linked to the mannequin and to the optical overlay, "
              "the visualization simultaneously communicates the result, its dynamics, and "
              "its validated accuracy.")

    # ---- 13 Results
    add_heading(doc, "13  Headline results (P01 / 2minWalk / right leg)", 1)
    table(doc, ["Joint", "RMSE vs optical (deg)", "ROM computed (deg)", "ROM optical (deg)"],
          [[j.capitalize(), f"{rmse[j]:.1f}",
            f"{rw[j]['computed']:.1f}", f"{rw[j]['optical']:.1f}"]
           for j in ["ankle", "knee", "hip"]])
    table(doc, ["Gait metric", "Value"],
          [["Walking bout", f"{summary['duration_s']:.1f} s continuous"],
           ["Cadence", f"{g['cadence_steps_per_min']:.0f} steps/min"],
           ["Stride time", f"{g['stride_time_mean_s']:.2f} ± {g['stride_time_std_s']:.2f} s"],
           ["Steady strides / strikes", f"{g['n_steady_strides']} / {g['n_foot_strikes']}"],
           ["Turnarounds", f"{len(summary['turnarounds'])}"],
           ["Heading RMSE 6-DOF / 9-DOF",
            f"{summary['heading_rmse_vs_optical_deg']['6dof']:.1f}° / "
            f"{summary['heading_rmse_vs_optical_deg']['9dof']:.1f}°"]])
    figure(doc, BASE + "_dashboard.png",
           "Figure 9. The full headless dashboard: per-joint computed vs optical flexion "
           "(left), angular velocity and foot impacts (right), and the metrics/heading panel.",
           width=6.6)

    # ---- 14 Caveats
    add_heading(doc, "14  Caveats (read before trusting numbers)", 1)
    for i, c in enumerate(summary["caveats"], 1):
        para(doc, f"{i}.  {c}")
    para(doc, "In short: 6-DOF is primary; the hip ROM is biased low (pelvis timing); use the "
              "in-window ROM, not the full-bout ROM; RMSE is offset-removed; and the "
              "magnetometer cannot help indoors here but the architecture supports it once "
              "field quality is gated on real hardware.", italic=True)

    # ---- 15 References
    add_heading(doc, "15  References", 1)
    refs = [
        "Madgwick, Harrison, Vaidyanathan. Estimation of IMU and MARG orientation using a "
        "gradient descent algorithm. IEEE ICORR, 2011.",
        "Zeni, Richards, Higginson. Two simple methods for determining gait events using "
        "kinematic data. Gait & Posture 27(4):710–714, 2008.",
        "Renaudin, Afzal, Lachapelle. Complete triaxis magnetometer calibration in the "
        "magnetic domain. J. Sensors, 2010.",
        "Seel, Raisch, Schauer. IMU-based joint angle measurement for gait analysis. "
        "Sensors 14(4):6891–6909, 2014.",
        "Grouvel et al., 2023 — the Geneva dataset (8 Physilog 6S IMUs + optical markers + "
        "force plates + pressure insoles).",
    ]
    for i, r in enumerate(refs, 1):
        para(doc, f"[{i}]  {r}", size=9.5)

    doc.save(DOCX)
    print(f"Wrote {DOCX}")


if __name__ == "__main__":
    build()
